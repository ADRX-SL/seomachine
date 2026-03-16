"""
Convert a markdown draft to a Word document with clean formatting.
Image lines ![alt](images/file.png) embed the actual PNG from drafts/images/ inline.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Match markdown image: ![anything](images/filename.png)
IMAGE_LINE_RE = re.compile(r'^!\[.*\]\(images/([^)]+)\)\s*$')


def replace_image_tags(text: str) -> str:
    """Replace ![alt](images/file.png) with [IMAGE: file.png] (for in-paragraph refs)."""
    return re.sub(r'!\[[^\]]*\]\(images/([^)]+)\)', r'[IMAGE: \1]', text)


def strip_bold(text: str) -> str:
    """Remove ** markers for plain text; caller can use bold in docx separately if needed."""
    return text  # Keep ** for now; we'll handle in add_paragraph


def add_paragraph_with_format(doc: Document, line: str, style_name: str = None):
    """Add a paragraph, handling **bold** and optional style."""
    line = replace_image_tags(line)
    # Simple bold handling: split by ** and add runs
    parts = re.split(r'(\*\*[^*]+\*\*)', line)
    p = doc.add_paragraph(style=style_name)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2] + ' ')
            run.bold = True
        else:
            if part:
                p.add_run(part)
    return p


def is_meta_line(line: str) -> bool:
    return (
        line.startswith('Meta title:') or
        line.startswith('Meta description:') or
        line.startswith('Target keyword:') or
        line.startswith('Slug:')
    )


def md_to_docx(md_path: Path, out_path: Path, embed_images: bool = True) -> None:
    doc = Document()
    content = md_path.read_text(encoding='utf-8')
    images_dir = md_path.parent / 'images'

    lines = content.splitlines()
    i = 0
    in_list = False
    list_style = None  # 'bullet' or 'number'

    while i < len(lines):
        line = lines[i]
        raw = line
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            in_list = False
            doc.add_paragraph()
            i += 1
            continue

        # Skip meta block
        if is_meta_line(stripped):
            i += 1
            continue

        # Skip HTML comments
        if stripped.startswith('<!--'):
            while i < len(lines) and '-->' not in lines[i]:
                i += 1
            i += 1
            continue

        # H1
        if stripped.startswith('# ') and not stripped.startswith('## '):
            in_list = False
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        # H2
        if stripped.startswith('## ') and not stripped.startswith('### '):
            in_list = False
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        # H3
        if stripped.startswith('### '):
            in_list = False
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        # Image line: ![alt](images/filename.png) — embed PNG or placeholder
        img_match = IMAGE_LINE_RE.match(stripped)
        if img_match:
            in_list = False
            filename = img_match.group(1)
            image_path = images_dir / filename
            if embed_images and image_path.exists():
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(str(image_path), width=Inches(5.5))
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(12)
            else:
                p = doc.add_paragraph(f'[IMAGE: {filename}]')
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Legacy placeholder line (if content was pre-processed)
        if stripped.startswith('[IMAGE:'):
            in_list = False
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            in_list = False
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:].strip())
            run.italic = True
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Numbered list (1. 2. etc.)
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            in_list = False
            _, rest = num_match.groups()
            p = doc.add_paragraph(style='List Number')
            # Handle bold in rest
            parts = re.split(r'(\*\*[^*]+\*\*)', rest)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    if part:
                        p.add_run(part)
            i += 1
            continue

        # Bullet list (- or *)
        if stripped.startswith('- ') or (stripped.startswith('* ') and len(stripped) > 2):
            rest = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*[^*]+\*\*)', rest)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    if part:
                        p.add_run(part)
            in_list = True
            i += 1
            continue

        # Continuation of list (indented with spaces)
        if line.startswith('  ') and ('- ' in line or '* ' in line.strip()[:3]):
            rest = line.strip().lstrip('-* ').strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*[^*]+\*\*)', rest)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    if part:
                        p.add_run(part)
            i += 1
            continue

        # Normal paragraph
        in_list = False
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                if part:
                    p.add_run(part)

        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python scripts/md_to_docx.py drafts/post.md [output.docx]", file=sys.stderr)
        sys.exit(1)
    md_path = Path(sys.argv[1]).resolve()
    if not md_path.exists():
        print(f"File not found: {md_path}", file=sys.stderr)
        sys.exit(1)
    out_path = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else (Path.cwd() / 'output' / (md_path.stem + '.docx')).resolve()
    try:
        md_to_docx(md_path, out_path)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)
